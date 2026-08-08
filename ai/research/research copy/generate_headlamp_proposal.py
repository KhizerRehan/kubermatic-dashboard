#!/usr/bin/env python3
"""Generate Headlamp Migration Proposal DOCX with diagrams and flow charts."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

# Monospace style for diagrams
mono_style = doc.styles.add_style('Mono', WD_STYLE_TYPE.PARAGRAPH)
mono_style.font.name = 'Courier New'
mono_style.font.size = Pt(8)
mono_style.paragraph_format.space_before = Pt(2)
mono_style.paragraph_format.space_after = Pt(0)
mono_style.paragraph_format.line_spacing = 1.0

mono_bold = doc.styles.add_style('MonoBold', WD_STYLE_TYPE.PARAGRAPH)
mono_bold.font.name = 'Courier New'
mono_bold.font.size = Pt(8)
mono_bold.font.bold = True
mono_bold.paragraph_format.space_before = Pt(2)
mono_bold.paragraph_format.space_after = Pt(0)
mono_bold.paragraph_format.line_spacing = 1.0


def add_mono_block(lines, bold_first=False):
    """Add a monospaced text block (for ASCII diagrams)."""
    for i, line in enumerate(lines):
        style_name = 'MonoBold' if (bold_first and i == 0) else 'Mono'
        doc.add_paragraph(line, style=style_name)


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    return table


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))


# ============================================================================
# TITLE PAGE
# ============================================================================

for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Migration Proposal:\nKubernetes Dashboard to Headlamp')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Kubermatic Kubernetes Platform (KKP) v2.31')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('')

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    'Issue: github.com/kubermatic/kubermatic/issues/15287\n'
    'Date: April 2026\n'
    'Status: Draft Proposal'
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS (manual)
# ============================================================================

doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1.', 'Executive Summary'),
    ('2.', 'Background & Motivation'),
    ('3.', 'Current Architecture (Before)'),
    ('4.', 'Proposed Architecture (After)'),
    ('5.', 'Architecture Comparison'),
    ('6.', 'Migration Strategy'),
    ('7.', 'API Changes'),
    ('8.', 'Implementation Plan'),
    ('9.', 'File Change Map'),
    ('10.', 'Risk Assessment'),
    ('11.', 'Pre-Implementation Blockers'),
    ('12.', 'Verification & Testing Plan'),
    ('13.', 'Timeline & Milestones'),
    ('14.', 'Open Questions'),
    ('15.', 'Appendix: Validated Server Flags'),
]
for num, title_text in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title_text}')
    run.font.size = Pt(12)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================

doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'The upstream Kubernetes Dashboard project (github.com/kubernetes-retired/dashboard) '
    'has been officially retired and archived. Continued use poses security risks due to '
    'lack of maintenance and security patches. This proposal outlines the migration to '
    'Headlamp (github.com/kubernetes-sigs/headlamp), a modern, extensible Kubernetes web '
    'UI maintained under the kubernetes-sigs organization.'
)

doc.add_heading('Key Decisions', level=2)
add_bullet('Deploy Headlamp on the seed cluster (per-user-cluster namespace), mirroring the current Kubernetes Dashboard pattern')
add_bullet('Add new Headlamp API field alongside the deprecated KubernetesDashboard field for backward compatibility')
add_bullet('Use existing cluster reconciler infrastructure (Go reconcilers), not the Application Framework')
add_bullet('Scope: Add Headlamp only; old dashboard code remains intact during migration period')

doc.add_heading('Requirements (from Issue #15287)', level=2)
add_bullet('Enable end-users to deploy and expose Headlamp through KKP')
add_bullet('Plugin support for: Cert Manager, KEDA, Flux (follow-up investigation)')
add_bullet('Migration strategy to remove old Kubernetes Dashboard')
add_bullet('OIDC integration support (follow-up investigation)')

# ============================================================================
# 2. BACKGROUND
# ============================================================================

doc.add_heading('2. Background & Motivation', level=1)

doc.add_heading('Why Migrate?', level=2)
add_bullet('Kubernetes Dashboard repository is archived (kubernetes-retired/dashboard)')
add_bullet('No security patches or maintenance since retirement')
add_bullet('Headlamp is actively maintained under kubernetes-sigs')
add_bullet('Headlamp provides modern UI, plugin system, and native OIDC support')
add_bullet('CNCF community backing ensures long-term viability')

doc.add_heading('Why Seed-Side Deployment (Not Application Framework)?', level=2)
doc.add_paragraph(
    'Earlier proposals (deprecated plans v1/v2) suggested deploying Headlamp via the KKP '
    'Application Framework using Helm charts installed directly into user clusters. After '
    'analysis, the seed-side approach was chosen for these reasons:'
)
add_bullet('Matches the proven, existing Kubernetes Dashboard deployment pattern')
add_bullet('No architectural change to how users access the dashboard')
add_bullet('Security: dashboard credentials stay on the seed cluster')
add_bullet('Simpler migration path: parallel deployment during transition')
add_bullet('No dependency on Application Framework for core platform components')

# ============================================================================
# 3. CURRENT ARCHITECTURE
# ============================================================================

doc.add_heading('3. Current Architecture (Before)', level=1)

doc.add_heading('3.1 Deployment Topology', level=2)
doc.add_paragraph(
    'The Kubernetes Dashboard is deployed as a split architecture: the main dashboard '
    'runs on the seed cluster, while supporting components (metrics-scraper, RBAC, secrets) '
    'are deployed in the user cluster.'
)

add_mono_block([
    '                    CURRENT ARCHITECTURE',
    '',
    '  SEED CLUSTER (per-user-cluster namespace: cluster-xyz)',
    '  +------------------------------------------------------+',
    '  |                                                      |',
    '  |  Deployment: kubernetes-dashboard (2 replicas)       |',
    '  |    Image: kubernetesui/dashboard:v2.7.0              |',
    '  |    Port: 9090                                        |',
    '  |    Command: /dashboard                               |',
    '  |                                                      |',
    '  |  Secret: kubernetes-dashboard-kubeconfig              |',
    '  |    Cert user: kubermatic:kubernetes-dashboard         |',
    '  |                                                      |',
    '  |  Seed Controller Manager:                            |',
    '  |    - DeploymentReconciler (creates dashboard)        |',
    '  |    - KubeconfigReconciler (creates kubeconfig)       |',
    '  |    - HealthCheck (monitors deployment health)        |',
    '  +-------------------+----------------------------------+',
    '                      | kubeconfig (cert auth)',
    '                      v',
    '  USER CLUSTER',
    '  +------------------------------------------------------+',
    '  |                                                      |',
    '  |  Namespace: kubernetes-dashboard                     |',
    '  |    +-- dashboard-metrics-scraper (2 replicas)        |',
    '  |    |   Image: kubernetesui/metrics-scraper:v1.0.8    |',
    '  |    |   Port: 8000                                    |',
    '  |    +-- Service: dashboard-metrics-scraper            |',
    '  |    +-- ServiceAccount                                |',
    '  |    +-- Role + RoleBinding                            |',
    '  |    +-- Secrets: key-holder, csrf                     |',
    '  |                                                      |',
    '  |  ClusterRole: system:dashboard-metrics-scraper       |',
    '  |  ClusterRoleBinding                                  |',
    '  |                                                      |',
    '  |  User Cluster Controller Manager:                    |',
    '  |    Reconciles all above resources                    |',
    '  +------------------------------------------------------+',
])

doc.add_heading('3.2 Control Flow', level=2)

add_mono_block([
    '                 Cluster Creation / Reconciliation',
    '                              |',
    '                              v',
    '              +-------------------------------+',
    '              | ClusterSpec.KubernetesDashboard|',
    '              |         .Enabled?              |',
    '              +---------------+---------------+',
    '                              |',
    '                    +---------+---------+',
    '                    |                   |',
    '                    v                   v',
    '               Enabled=true        Enabled=false',
    '                    |                   |',
    '           +--------+--------+          |',
    '           |                 |          v',
    '           v                 v     Delete All',
    '    Seed Controller    User Cluster  Resources',
    '    Manager            Controller',
    '           |                 |',
    '           v                 v',
    '    - Deploy dashboard  - Create namespace',
    '      deployment          (kubernetes-dashboard)',
    '    - Create kubeconfig - Deploy metrics-scraper',
    '      secret            - Create RBAC',
    '    - Health check      - Create secrets',
    '                        - Create service',
])

doc.add_heading('3.3 Current File Structure', level=2)

add_mono_block([
    '  pkg/',
    '  +-- resources/',
    '  |   +-- kubernetes-dashboard/          <-- SEED-SIDE',
    '  |       +-- deployment.go              |   Dashboard deployment reconciler',
    '  |       +-- deletion.go                |   Cleanup when disabled',
    '  |',
    '  +-- controller/',
    '  |   +-- seed-controller-manager/',
    '  |   |   +-- kubernetes/',
    '  |   |       +-- resources.go           |   Wires dashboard reconcilers',
    '  |   |       +-- health.go              |   Dashboard health check',
    '  |   |',
    '  |   +-- user-cluster-controller-manager/',
    '  |       +-- resources/',
    '  |           +-- reconciler.go          |   Wires user-cluster reconcilers',
    '  |           +-- resources/',
    '  |               +-- kubernetes-dashboard/  <-- USER-CLUSTER',
    '  |                   +-- constants.go       |   Namespace, AppName',
    '  |                   +-- namespace.go       |   Namespace with PSA labels',
    '  |                   +-- deployment.go      |   Metrics-scraper',
    '  |                   +-- service.go         |   Metrics-scraper service',
    '  |                   +-- serviceaccount.go  |   SA',
    '  |                   +-- role.go            |   RBAC role',
    '  |                   +-- rolebinding.go     |   RBAC binding',
    '  |                   +-- clusterrole.go     |   Metrics ClusterRole',
    '  |                   +-- clusterrolebinding.go',
    '  |                   +-- secret.go          |   Key-holder + CSRF',
    '  |                   +-- deletion.go        |   Cleanup list',
    '  |',
    '  +-- resources/resources.go             |   Resource name constants',
    '  +-- defaulting/cluster.go              |   Default: enabled=true',
    '  +-- install/images/images.go           |   Image collection',
    '  +-- resources/test/fixtures/           |   ~66 deployment fixtures',
])

# ============================================================================
# 4. PROPOSED ARCHITECTURE
# ============================================================================

doc.add_heading('4. Proposed Architecture (After)', level=1)

doc.add_heading('4.1 Headlamp Deployment Topology', level=2)
doc.add_paragraph(
    'Headlamp is deployed on the seed cluster in the per-user-cluster namespace, using a '
    'kubeconfig secret to connect to the user cluster API server. This mirrors the exact '
    'pattern of the current Kubernetes Dashboard but with a dramatically simpler user-cluster '
    'footprint (no metrics-scraper, no secrets, no services).'
)

add_mono_block([
    '                    PROPOSED ARCHITECTURE',
    '',
    '  SEED CLUSTER (per-user-cluster namespace: cluster-xyz)',
    '  +------------------------------------------------------+',
    '  |                                                      |',
    '  |  Deployment: headlamp (2 replicas)                   |',
    '  |    Image: ghcr.io/headlamp-k8s/headlamp:v0.26.0     |',
    '  |    Command: /headlamp/headlamp-server                |',
    '  |    Args: -kubeconfig /etc/kubernetes/kubeconfig       |',
    '  |          -html-static-dir /headlamp/frontend         |',
    '  |    Port: 4466                                        |',
    '  |    Security: runAsUser 1001, readOnlyRootFilesystem  |',
    '  |    Resources: 100m/128Mi request, 250m/256Mi limit   |',
    '  |                                                      |',
    '  |  Secret: headlamp-kubeconfig                         |',
    '  |    (cert user: kubermatic:headlamp)                  |',
    '  |    Generated via GetInternalKubeconfigReconciler      |',
    '  |                                                      |',
    '  |  Seed Controller Manager:                            |',
    '  |    - DeploymentReconciler (creates headlamp)         |',
    '  |    - KubeconfigReconciler (creates kubeconfig)       |',
    '  |    - HealthCheck (monitors deployment health)        |',
    '  +-------------------+----------------------------------+',
    '                      | kubeconfig -> user cluster API',
    '                      v',
    '  USER CLUSTER',
    '  +------------------------------------------------------+',
    '  |                                                      |',
    '  |  Namespace: headlamp (PSA baseline labels)           |',
    '  |                                                      |',
    '  |  ClusterRole: system:headlamp                        |',
    '  |    - Read access to cluster-scoped resources          |',
    '  |    - (namespaces, nodes, CRDs, workloads, etc.)      |',
    '  |                                                      |',
    '  |  ClusterRoleBinding: system:headlamp                 |',
    '  |    Subject: User "kubermatic:headlamp"               |',
    '  |    RoleRef: ClusterRole "system:headlamp"            |',
    '  |                                                      |',
    '  |  NO deployments, services, or secrets                |',
    '  |  (Headlamp reads metrics API directly)               |',
    '  +------------------------------------------------------+',
])

doc.add_heading('4.2 How Users Access Headlamp', level=2)

add_mono_block([
    '  Access Flow:',
    '',
    '    End User',
    '       |',
    '       | (KKP Dashboard UI exposes proxy URL,',
    '       |  same pattern as current k8s-dashboard)',
    '       v',
    '    KKP API Server / Proxy',
    '       |',
    '       | route to seed cluster',
    '       v',
    '    Headlamp Service (seed, cluster-xyz ns)',
    '       |',
    '       | port 4466',
    '       v',
    '    Headlamp Pod',
    '       |',
    '       | uses kubeconfig secret',
    '       v',
    '    User Cluster API Server',
    '       |',
    '       v',
    '    Cluster Resources (namespaces, pods, etc.)',
])

# ============================================================================
# 5. COMPARISON
# ============================================================================

doc.add_heading('5. Architecture Comparison', level=1)

doc.add_heading('5.1 Before vs After', level=2)

add_table(
    ['Aspect', 'BEFORE (Kubernetes Dashboard)', 'AFTER (Headlamp)'],
    [
        ['Tool', 'kubernetes-dashboard v2.7.0 (archived)', 'Headlamp v0.26.0 (active, kubernetes-sigs)'],
        ['Seed Deployment', 'kubernetes-dashboard (2 replicas)', 'headlamp (2 replicas)'],
        ['Container Port', '9090', '4466'],
        ['User-cluster Deployments', 'dashboard-metrics-scraper', 'None'],
        ['User-cluster RBAC', 'Role + ClusterRole (metrics only)', 'ClusterRole only (full browsing)'],
        ['User-cluster Secrets', 'JWE key-holder + CSRF token', 'None'],
        ['User-cluster ServiceAccount', 'dashboard-metrics-scraper', 'None'],
        ['User-cluster Service', 'dashboard-metrics-scraper:8000', 'None'],
        ['Container Image', 'kubernetesui/dashboard', 'ghcr.io/headlamp-k8s/headlamp'],
        ['Plugin Support', 'Not possible', 'Native plugin system'],
        ['OIDC Support', 'Not implemented', 'Native support (follow-up)'],
        ['Health Check', 'ExtendedHealth.KubernetesDashboard', 'ExtendedHealth.Headlamp'],
        ['Go files (seed)', '2 files', '2 files (same pattern)'],
        ['Go files (user cluster)', '11 files', '4 files (much simpler)'],
    ]
)

doc.add_heading('5.2 Why Headlamp is Simpler on User-Cluster Side', level=2)
add_bullet('Built-in resource browsing (no metrics-scraper sidecar needed)')
add_bullet('Handles auth via kubeconfig (no JWE/CSRF secrets needed)')
add_bullet('Needs cluster-wide read access (ClusterRole) since it is a full cluster browser')
add_bullet('No in-cluster service, deployment, or service account required')

# ============================================================================
# 6. MIGRATION STRATEGY
# ============================================================================

doc.add_heading('6. Migration Strategy', level=1)

doc.add_heading('6.1 Phased Approach', level=2)
doc.add_paragraph(
    'The migration uses a parallel-deployment strategy: Headlamp is added alongside the '
    'existing Kubernetes Dashboard. Both can run simultaneously during the transition period. '
    'The old dashboard code is NOT removed in this scope.'
)

add_mono_block([
    '  Migration Timeline',
    '',
    '  KKP v2.30 (current)     KKP v2.31 (this work)    KKP v2.32+ (future)',
    '  +-------------------+   +-------------------+    +-------------------+',
    '  |                   |   |                   |    |                   |',
    '  | k8s-dashboard     |   | k8s-dashboard     |    |                   |',
    '  | (deployed)        |   | (still deployed)  |    | (REMOVED)         |',
    '  |                   |   |                   |    |                   |',
    '  |                   |   | + Headlamp        |    | Headlamp          |',
    '  |                   |   |   (NEW, deployed)  |    | (sole dashboard)  |',
    '  |                   |   |                   |    |                   |',
    '  | API field:        |   | API fields:       |    | API field:        |',
    '  | kubernetesDashboard|   | kubernetesDashboard|    | headlamp          |',
    '  |                   |   | + headlamp (NEW)  |    | (old removed)     |',
    '  +-------------------+   +-------------------+    +-------------------+',
])

doc.add_heading('6.2 CRD Field Migration Flow', level=2)

add_mono_block([
    '  Defaulting Webhook Logic (KKP v2.31)',
    '',
    '  +-- Is spec.headlamp set? --+',
    '  |                           |',
    '  NO                         YES',
    '  |                           |',
    '  v                           v',
    '  Set spec.headlamp =        Use as-is',
    '  { enabled: true }          (user explicitly configured)',
    '  (default for all clusters)',
    '',
    '  Note: spec.kubernetesDashboard remains untouched.',
    '  Both fields coexist. Old field deprecated in v2.32+.',
])

doc.add_heading('6.3 IsHeadlampEnabled() Logic', level=2)

add_mono_block([
    '  func IsHeadlampEnabled() bool:',
    '',
    '    +-- spec.Headlamp != nil? --+',
    '    |                           |',
    '   YES                         NO',
    '    |                           |',
    '    v                           v',
    '  return                    return TRUE',
    '  spec.Headlamp.Enabled     (default: enabled)',
    '',
    '  Simple: if the field exists, use it.',
    '  If it does not exist (nil), default to enabled.',
])

# ============================================================================
# 7. API CHANGES
# ============================================================================

doc.add_heading('7. API Changes', level=1)

doc.add_heading('7.1 New Types (sdk/apis/kubermatic/v1/cluster.go)', level=2)

add_mono_block([
    '  // Headlamp contains settings for the Headlamp component.',
    '  type Headlamp struct {',
    '      // Controls whether Headlamp is deployed.',
    '      // Enabled by default.',
    '      Enabled bool `json:"enabled,omitempty"`',
    '  }',
    '',
    '  // Added to ClusterSpec:',
    '  Headlamp *Headlamp `json:"headlamp,omitempty"`',
    '',
    '  // New method:',
    '  func (c ClusterSpec) IsHeadlampEnabled() bool {',
    '      return c.Headlamp == nil || c.Headlamp.Enabled',
    '  }',
])

doc.add_heading('7.2 Health Status (cluster_status.go)', level=2)

add_mono_block([
    '  // Added to ExtendedClusterHealth:',
    '  Headlamp *HealthStatus `json:"headlamp,omitempty"`',
])

doc.add_heading('7.3 New Constants (pkg/resources/resources.go)', level=2)

add_table(
    ['Constant', 'Value'],
    [
        ['HeadlampDeploymentName', '"headlamp"'],
        ['HeadlampKubeconfigSecretName', '"headlamp-kubeconfig"'],
        ['HeadlampCertUsername', '"kubermatic:headlamp"'],
        ['HeadlampClusterRoleName', '"system:headlamp"'],
        ['HeadlampClusterRoleBindingName', '"system:headlamp"'],
    ]
)

doc.add_heading('7.4 Backward Compatibility', level=2)
doc.add_paragraph(
    'The existing KubernetesDashboard field and IsKubernetesDashboardEnabled() method '
    'remain unchanged. They are NOT removed in this scope. Both dashboards can coexist. '
    'Deprecation and removal of the old field will be handled in KKP v2.32+.'
)

# ============================================================================
# 8. IMPLEMENTATION PLAN
# ============================================================================

doc.add_heading('8. Implementation Plan', level=1)

doc.add_heading('Phase 1: API Types & Constants', level=2)
add_bullet('Add Headlamp struct and field to ClusterSpec')
add_bullet('Add IsHeadlampEnabled() method')
add_bullet('Add Headlamp health status to ExtendedClusterHealth')
add_bullet('Add new constants to resources.go')
add_bullet('Add defaulting logic in cluster.go')

doc.add_heading('Phase 2: Seed-Side Resources', level=2)
add_bullet('Create pkg/resources/headlamp/deployment.go (deployment reconciler)')
add_bullet('Create pkg/resources/headlamp/deletion.go (cleanup resources)')
add_bullet('Wire DeploymentReconciler into seed controller resources.go')
add_bullet('Wire kubeconfig secret reconciler into seed controller')
add_bullet('Add headlampHealthCheck to health.go')
add_bullet('Add cleanup when IsHeadlampEnabled() == false')

doc.add_heading('Phase 3: User-Cluster Resources', level=2)
add_bullet('Create pkg/controller/.../headlamp/constants.go')
add_bullet('Create pkg/controller/.../headlamp/namespace.go (PSA baseline)')
add_bullet('Create pkg/controller/.../headlamp/clusterrole.go (read access)')
add_bullet('Create pkg/controller/.../headlamp/clusterrolebinding.go')
add_bullet('Create pkg/controller/.../headlamp/deletion.go')
add_bullet('Wire into user-cluster reconciler.go')

doc.add_heading('Phase 4: Image Registration & Versions', level=2)
add_bullet('Add headlamp image to pkg/install/images/images.go')
add_bullet('Add headlamp entry to hack/versions.yaml')

doc.add_heading('Phase 5: Testing & Verification', level=2)
add_bullet('Add deployment reconciler golden test fixtures')
add_bullet('Run CRD generation (make generate)')
add_bullet('Manual testing: create cluster, verify deployment')
add_bullet('Disable testing: set enabled=false, verify cleanup')
add_bullet('Coexistence testing: both dashboards running simultaneously')

# ============================================================================
# 9. FILE CHANGE MAP
# ============================================================================

doc.add_heading('9. File Change Map', level=1)

doc.add_heading('9.1 New Files to Create', level=2)

add_mono_block([
    '  NEW FILES:',
    '  +-- pkg/resources/headlamp/',
    '  |   +-- deployment.go              Headlamp deployment reconciler',
    '  |   +-- deletion.go                Seed-side cleanup resources',
    '  |',
    '  +-- pkg/controller/user-cluster-controller-manager/',
    '      +-- resources/resources/headlamp/',
    '          +-- constants.go            Namespace & AppName constants',
    '          +-- namespace.go            Headlamp namespace with PSA',
    '          +-- clusterrole.go          Read-access ClusterRole',
    '          +-- clusterrolebinding.go   Bind to kubermatic:headlamp user',
    '          +-- deletion.go             User-cluster cleanup resources',
])

doc.add_heading('9.2 Files to Modify', level=2)

add_table(
    ['File', 'Change'],
    [
        ['sdk/apis/kubermatic/v1/cluster.go', 'Add Headlamp struct, field, IsHeadlampEnabled()'],
        ['sdk/apis/kubermatic/v1/cluster_status.go', 'Add Headlamp to ExtendedClusterHealth'],
        ['pkg/resources/resources.go', 'Add Headlamp constants'],
        ['pkg/defaulting/cluster.go', 'Add Headlamp defaulting (enabled=true)'],
        ['pkg/controller/seed-controller-manager/kubernetes/resources.go', 'Wire deployment + kubeconfig reconcilers'],
        ['pkg/controller/seed-controller-manager/kubernetes/health.go', 'Add headlampHealthCheck'],
        ['pkg/controller/user-cluster-controller-manager/resources/reconciler.go', 'Wire namespace, ClusterRole, ClusterRoleBinding, cleanup'],
        ['pkg/install/images/images.go', 'Add headlamp image for mirroring'],
        ['hack/versions.yaml', 'Add headlamp version entry'],
    ]
)

doc.add_heading('9.3 Auto-Generated Files (via make generate)', level=2)
add_bullet('pkg/crd/k8c.io/kubermatic.k8c.io_clusters.yaml')
add_bullet('pkg/crd/k8c.io/kubermatic.k8c.io_clustertemplates.yaml')
add_bullet('sdk/apis/kubermatic/v1/zz_generated.deepcopy.go')

doc.add_heading('9.4 Files NOT Changed (Old Dashboard Intact)', level=2)
add_bullet('pkg/resources/kubernetes-dashboard/ -- remains as-is')
add_bullet('pkg/controller/.../kubernetes-dashboard/ -- remains as-is')
add_bullet('All existing k8s-dashboard test fixtures -- remain as-is')

# ============================================================================
# 10. RISK ASSESSMENT
# ============================================================================

doc.add_heading('10. Risk Assessment', level=1)

add_table(
    ['#', 'Risk', 'Impact', 'Likelihood', 'Mitigation'],
    [
        ['1', 'ClusterRole grants broad read access', 'HIGH', 'CERTAIN',
         'Review RBAC rules; restrict if needed'],
        ['2', 'Headlamp image compatibility (UID/GID)', 'MEDIUM', 'MEDIUM',
         'Helm defaults UID 100; spec uses 1001. Test both'],
        ['3', 'Plugin support not in initial scope', 'MEDIUM', 'CERTAIN',
         'Documented as follow-up; plugins need investigation'],
        ['4', 'OIDC integration not in initial scope', 'MEDIUM', 'CERTAIN',
         'Documented as follow-up; native support exists'],
        ['5', 'Health check monitoring gap during transition', 'LOW', 'LOW',
         'Both health checks run in parallel'],
        ['6', 'Image not mirrored for air-gap environments', 'HIGH', 'CERTAIN',
         'Add to quay.io/kubermatic-mirror pipeline'],
        ['7', 'Frontend (KKP Dashboard) needs updates', 'MEDIUM', 'CERTAIN',
         'Coordinate with frontend team for v2.31'],
        ['8', 'Coexistence issues (both dashboards)', 'LOW', 'LOW',
         'Different namespaces, no resource conflicts'],
    ]
)

doc.add_heading('10.1 Risk Matrix Visualization', level=2)

add_mono_block([
    '                            RISK MATRIX',
    '  HIGH    |                                              |',
    '  IMPACT  |  [6] Air-gap       [1] RBAC scope           |',
    '          |  mirroring                                   |',
    '          |                                              |',
    '  MEDIUM  |  [3] Plugins       [2] Image UID/GID        |',
    '  IMPACT  |  [4] OIDC          [7] Frontend updates     |',
    '          |                                              |',
    '  LOW     |                    [8] Coexistence           |',
    '  IMPACT  |                    [5] Health gap            |',
    '          +----------------------------------------------+',
    '            LOW LIKELIHOOD      MEDIUM          HIGH',
])

# ============================================================================
# 11. PRE-IMPLEMENTATION BLOCKERS
# ============================================================================

doc.add_heading('11. Pre-Implementation Blockers', level=1)

add_table(
    ['#', 'Item', 'Priority', 'Status'],
    [
        ['1', 'Verify headlamp container image works with UID 1001', 'BLOCKER', 'Needs testing'],
        ['2', 'Add headlamp image to quay.io/kubermatic-mirror', 'HIGH', 'Not started'],
        ['3', 'Define exact ClusterRole RBAC rules', 'HIGH', 'Draft in design spec'],
        ['4', 'Frontend team coordination for API field', 'MEDIUM', 'Not started'],
        ['5', 'Plugin investigation (Cert Manager, KEDA, Flux)', 'FOLLOW-UP', 'Deferred'],
        ['6', 'OIDC integration investigation', 'FOLLOW-UP', 'Deferred'],
    ]
)

# ============================================================================
# 12. VERIFICATION PLAN
# ============================================================================

doc.add_heading('12. Verification & Testing Plan', level=1)

doc.add_heading('12.1 Testing Flow', level=2)

add_mono_block([
    '  Testing Pipeline',
    '',
    '  +------------------+     +------------------+     +------------------+',
    '  |  Unit Tests      |     |  CRD Generation  |     |  Manual Testing  |',
    '  |                  |---->|                  |---->|                  |',
    '  |  go test         |     |  make generate   |     |  Deploy cluster  |',
    '  |  ./pkg/resources |     |  Verify CRD YAML |     |  Check headlamp  |',
    '  |  ./pkg/controller|     |  Check deepcopy  |     |  Check health    |',
    '  |  ./sdk/...       |     |                  |     |  Check RBAC      |',
    '  +------------------+     +------------------+     +------------------+',
    '                                                           |',
    '                                                           v',
    '  +------------------+     +------------------+     +------------------+',
    '  | Coexistence      |     | Disable Testing  |     | Upgrade Testing  |',
    '  |                  |<----|                  |<----|                  |',
    '  | Both dashboards  |     | Set enabled=false|     | Existing cluster |',
    '  | run side-by-side |     | Verify cleanup   |     | gets headlamp    |',
    '  | No conflicts     |     | No orphans       |     | k8s-dash intact  |',
    '  +------------------+     +------------------+     +------------------+',
])

doc.add_heading('12.2 Test Checklist', level=2)
add_bullet('Unit: go test ./pkg/resources/... ./pkg/controller/... ./sdk/...')
add_bullet('CRD: make generate produces valid CRD YAML with headlamp field')
add_bullet('Golden tests: New deployment fixtures for headlamp added')
add_bullet('Create cluster: headlamp deployment appears in seed namespace')
add_bullet('Kubeconfig: headlamp-kubeconfig secret is created')
add_bullet('User cluster: ClusterRole and ClusterRoleBinding exist')
add_bullet('Health: ExtendedHealth.Headlamp reports healthy')
add_bullet('Disable: spec.headlamp.enabled=false removes all resources')
add_bullet('Coexistence: Both kubernetes-dashboard and headlamp run simultaneously')
add_bullet('Air-gap: headlamp image available in mirrored registry')

# ============================================================================
# 13. TIMELINE
# ============================================================================

doc.add_heading('13. Timeline & Milestones', level=1)

add_mono_block([
    '  Implementation Timeline',
    '',
    '  Phase 1: API Types & Constants           [  ####  ]  ~1 day',
    '  Phase 2: Seed-Side Resources             [  ########  ]  ~2 days',
    '  Phase 3: User-Cluster Resources          [  ######  ]  ~1.5 days',
    '  Phase 4: Image & Version Registration    [  ##  ]  ~0.5 day',
    '  Phase 5: Testing & Verification          [  ########  ]  ~2 days',
    '  ---------------------------------------------------',
    '  Total estimated effort:                  ~7 working days',
    '',
    '  Target: KKP v2.31',
    '',
    '  Future (v2.32+):',
    '    - Remove old kubernetes-dashboard code',
    '    - Remove deprecated kubernetesDashboard API field',
    '    - Add plugin support (Cert Manager, KEDA, Flux)',
    '    - Add OIDC integration',
    '    - Update KKP Dashboard frontend',
])

# ============================================================================
# 14. OPEN QUESTIONS
# ============================================================================

doc.add_heading('14. Open Questions', level=1)

add_table(
    ['#', 'Question', 'Impact', 'Notes'],
    [
        ['1', 'UID: Helm chart defaults to UID 100/GID 101, but design uses 1001 (matching k8s-dashboard). Which does the container support?', 'HIGH', 'Needs testing before implementation'],
        ['2', 'Plugin support: How to configure Cert Manager, KEDA, Flux plugins in seed-side deployment?', 'MEDIUM', 'Deferred to follow-up'],
        ['3', 'OIDC: How does Headlamp OIDC integrate with KKP OIDC setup?', 'MEDIUM', 'Deferred to follow-up'],
        ['4', '-insecure-ssl flag: Needed for self-signed API server certs?', 'LOW', 'Test with actual clusters'],
        ['5', 'Base URL: Does Headlamp need -base-url for KKP proxy routing?', 'MEDIUM', 'Test with KKP proxy'],
    ]
)

# ============================================================================
# 15. APPENDIX
# ============================================================================

doc.add_heading('15. Appendix: Validated Server Flags', level=1)

doc.add_paragraph(
    'The following flags were validated against a live dev cluster on 2026-03-31 using '
    'Docker with a mounted kubeconfig.'
)

add_table(
    ['Flag', 'Default', 'Our Usage'],
    [
        ['-kubeconfig', '""', '/etc/kubernetes/kubeconfig/kubeconfig'],
        ['-in-cluster', 'false', 'Leave as default (do not set)'],
        ['-port', '4466', 'Use default'],
        ['-plugins-dir', '~/.config/Headlamp/plugins', '/headlamp/plugins'],
        ['-html-static-dir', '""', '/headlamp/frontend'],
        ['-base-url', '""', 'May set for routing'],
        ['-insecure-ssl', 'false', 'May need for self-signed certs'],
    ]
)

doc.add_heading('Validated Test Results', level=2)
add_bullet('Starts with -kubeconfig pointing to external cluster')
add_bullet('Proxy setup to external cluster API server works')
add_bullet('34 namespaces, 12 nodes, deployments all browsable via API proxy')
add_bullet('No auth prompt (kubeconfig provides authentication)')
add_bullet('Port 4466 works out of the box')
add_bullet('No -in-cluster flag needed (defaults to false)')

doc.add_heading('Container Details', level=2)
add_bullet('Binary: /headlamp/headlamp-server')
add_bullet('Image: ghcr.io/headlamp-k8s/headlamp:v0.26.0')
add_bullet('Static dir: /headlamp/frontend (in container)')
add_bullet('Plugins dir: /headlamp/plugins (in container)')
add_bullet('Default RunAs: user headlamp (UID 100, GID 101 per Helm chart)')

# ============================================================================
# SAVE
# ============================================================================

output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'Headlamp-Migration-Proposal-KKP-v2.31.docx'
)
doc.save(output_path)
print(f'Saved: {output_path}')
